from __future__ import annotations

import json
from pathlib import Path

import pytest

from jobhunt.errors import PipelineError
from jobhunt.resume.parse_docx import (
    _ROLE_LINE_RE,
    _split_skills,
    parse_baseline,
    write_kb_markdown,
    write_verified_json,
)

REPO_ROOT = Path(__file__).resolve().parent.parent


def _live_baseline() -> Path | None:
    """The real baseline resume, or None when the repo has none.

    Resolved through `resume.locate` rather than a hard-coded filename: on
    2026-07-24 the file was renamed and the previous `REPO_ROOT /
    "Baseline_Resume.docx"` turned two regression guards into silent skips,
    which is strictly worse than a failure because nothing surfaces it.

    These are the **only** tests that touch the user's real resume, and they
    assert it parses cleanly rather than asserting its content. Content
    assertions belong on the fictional fixture profile (tests/conftest.py).
    """
    from jobhunt.errors import PipelineError
    from jobhunt.resume.locate import find_baseline_resume

    try:
        found = find_baseline_resume(REPO_ROOT)
    except PipelineError:
        return None
    return found if found.suffix.lower() == ".docx" else None


BASELINE = _live_baseline() or REPO_ROOT / "Baseline_Resume.docx"


def test_contact_block_spans_multiple_paragraphs(tmp_path: Path):
    """A contact block wrapped across two paragraphs is joined into one line."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com  |")
    doc.add_paragraph("https://janedev.com  |  https://github.com/jane")
    doc.add_paragraph("SUMMARY")
    doc.add_paragraph("A developer.")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, _ = parse_baseline(path)
    assert facts.name == "Jane Dev"
    assert "jane@example.com" in facts.contact_line
    assert "janedev.com" in facts.contact_line  # second paragraph captured
    assert "github.com/jane" in facts.contact_line
    assert facts.summary == "A developer."  # contact block didn't swallow the summary


@pytest.mark.skipif(not BASELINE.is_file(), reason="baseline .docx not present")
def test_parse_baseline_round_trip(tmp_path: Path):
    facts, warnings = parse_baseline(BASELINE)
    assert warnings == []  # the curated master parses cleanly (RP1 regression guard)
    assert facts.name
    assert "Toronto" in facts.contact_line
    # The master wraps contact info onto a second paragraph (links on line 2);
    # the whole block must be captured, not just the first line.
    assert "github.com/SimBuds" in facts.contact_line
    assert "caseyhsu.com" in facts.contact_line
    assert len(facts.work_history) == 4

    employers = {r.employer for r in facts.work_history}
    # Substring, not equality: the employer cell carries whatever detail the
    # current resume puts beside the name (a descriptor in one draft, a location
    # in the next). Pinning the exact string makes an intentional resume edit
    # look like a parser regression — see IMPLEMENT.md Phase A12.
    assert any("Atelier Dacko" in e for e in employers)
    # Every employer cell must be clean of stray separators: `employer` is half
    # the identity key the fabrication guard compares on.
    for e in employers:
        assert not e.endswith(("|", "-", "—", ",")), f"stray separator: {e!r}"
    assert "Sous Chef & Team Lead" in {r.title for r in facts.work_history}

    # Familiar must stay separate from Core when populated. The 2026-07-27 resume
    # drops the familiar tier entirely (it previously carried Java/Spring Boot as
    # coursework-only), so this asserts the disjointness property rather than
    # pinning specific tools the resume may no longer list.
    assert not set(facts.skills_familiar) & set(facts.skills_core)
    # Python is production experience and must never land in Familiar. Which
    # Core-side bucket it occupies depends on how the resume groups its rows,
    # so assert the honesty property rather than the bucket (Phase A12).
    assert "Python" not in facts.skills_familiar
    assert any(
        "Python" in s
        for s in facts.skills_core + facts.skills_data_devops + facts.skills_ai
    )

    # PB1: when the resume carries a "Project Stack:" row, it populates a
    # Core-grade bucket distinct from Familiar (project-demonstrated, not
    # academic). Conditional because not every draft includes that row.
    for item in facts.skills_projects:
        assert item not in facts.skills_familiar

    # PB3: the PROJECTS narrative section parses into structured projects.
    # Asserted structurally: how many projects the author keeps on the resume is
    # an editorial decision that changes between drafts, but every project that
    # IS present must come through with a clean name, a scheme-stripped url, and
    # its bullets attached. Named-project coverage belongs on a fixture, not the
    # live document — see IMPLEMENT.md Phase A12.
    assert facts.projects, "PROJECTS section produced no structured projects"
    for p in facts.projects:
        assert p.name and "|" not in p.name, f"bad project name: {p.name!r}"
        assert p.url and not p.url.startswith(("http://", "https://"))
        assert p.bullets, f"project {p.name!r} parsed with no bullets"
    # A "Stack:" line under a project populates that project's stack rather than
    # being swallowed as a bullet — assert both halves, since the failure mode is
    # the row landing in bullets rather than disappearing outright.
    assert any(p.stack for p in facts.projects), "no project captured a Stack: line"
    for p in facts.projects:
        assert not any(b.startswith("Stack:") for b in p.bullets), (
            f"project {p.name!r} swallowed a Stack: row as a bullet"
        )
    # PROJECTS narrative must NOT leak into education (the pre-PB3 behaviour).
    assert not any("github.com" in e for e in facts.education)

    # Round-trip via verified.json.
    out = tmp_path / "verified.json"
    write_verified_json(facts, out)
    payload = json.loads(out.read_text())
    assert payload["name"] == facts.name
    assert len(payload["work_history"]) == len(facts.work_history)
    assert len(payload["projects"]) == len(facts.projects)
    # Nested dataclass round-trips. Asserted by equality rather than truthiness so
    # a resume whose projects carry no Stack: row still exercises the round-trip.
    assert payload["projects"][0]["stack"] == facts.projects[0].stack

    # KB markdown writer leaves five files (projects.md added when projects exist).
    kb = tmp_path / "kb"
    paths = write_kb_markdown(facts, kb)
    assert len(paths) == 5
    skills_md = (kb / "profile" / "skills.md").read_text()
    assert skills_md.count("## Familiar") == 1
    assert skills_md.count("## Project Stack") == 1
    assert "FastAPI" in skills_md
    projects_md = (kb / "profile" / "projects.md").read_text()
    # Every parsed project gets its own H2 in the markdown sidecar. Asserted
    # against what the resume actually carries rather than a fixed project name,
    # so trimming the PROJECTS section is not a test failure (Phase A12).
    for p in facts.projects:
        assert f"## {p.name}" in projects_md


@pytest.mark.skipif(not BASELINE.is_file(), reason="baseline .docx not present")
def test_parse_baseline_positioning_and_atomic_skills():
    """Regression guard for the 2026-06 specialist re-positioning (Initiative:
    Positioning resync). Locks the things that were either hand-patched or newly
    moved, so a future re-parse can't silently regress them:
    - skills_ai stays ATOMIC and paren-aware (retires PLAN.md's stale "skills_ai
      produces a run-on, patch by hand" caveat).
    - Figma sits in Core as of the 2026-07-27 resume (superseding the 2026-06-18
      call that kept it Familiar). Tier is the author's positioning decision and
      is driven by which row the resume lists it under, not by this test.
    - "Dawn" survives the parse (it lives in the Atelier bullet, not the skills
      line).
    - The lead role carries the JD-aligned retitle, not the old generic title.
    """
    facts, warnings = parse_baseline(BASELINE)
    assert warnings == []

    # Atomic: no item is a comma-split artifact (a stray dangling close-paren
    # with no open), which is what a naive comma split produces on an entry like
    # "local LLM hosting and inference tuning (Ollama)".
    #
    # Requiring at least one PARENTHESISED item used to live here too, but that
    # was content-coupled despite the comment claiming otherwise: the 2026-08-10
    # rewrite legitimately dropped every parenthetical from the AI row
    # ("Claude API, Ollama, Model Context Protocol, ...") and the assertion
    # started failing on a resume edit, exactly the false positive it was
    # supposed to avoid. Paren-awareness itself is covered content-independently
    # by `test_split_skills_paren_aware`, so only the invariant belongs here.
    for item in facts.skills_ai:
        assert item.count("(") == item.count(")"), f"unbalanced parens: {item!r}"

    # Familiar and Core must stay disjoint — that split is the hard honesty
    # signal (`kb/policies/tailoring-rules.md`). Whether Familiar is populated at
    # all is the author's call: the 2026-07-27 resume drops the familiar tier
    # entirely, so emptiness is a valid state and only the disjointness holds.
    assert not set(facts.skills_familiar) & set(facts.skills_core)

    # Dawn is captured somewhere in the verified facts (it lives in a bullet).
    assert any("Dawn" in b for r in facts.work_history for b in r.bullets)

    # The lead role carries a real title and is the current ("Present") one.
    # The exact wording is the author's positioning call and changes between
    # drafts; what must hold is that a title parsed at all (a title-less header
    # is what the dash-form regression produced) — see IMPLEMENT.md Phase A12.
    lead = facts.work_history[0]
    assert lead.title, "lead role parsed with an empty title"
    assert "Present" in lead.dates


def test_parse_warns_on_unknown_skill_label(tmp_path: Path):
    """An unrecognized TECHNICAL SKILLS label is reported as a warning, not
    silently dropped (RP1 + RP3). 'Hobbies' is neither a canonical bucket nor an
    RP3 alias, so it stays unknown."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Core: Python, TypeScript")
    doc.add_paragraph("Hobbies: Climbing, Chess")  # unknown label, no bucket / alias
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert "Python" in facts.skills_core
    assert any("Hobbies" in w for w in warnings)
    # The dropped items must not have leaked into any bucket.
    all_skills = (
        facts.skills_core
        + facts.skills_cms
        + facts.skills_data_devops
        + facts.skills_ai
        + facts.skills_projects
        + facts.skills_familiar
    )
    assert "Climbing" not in all_skills


def test_skill_label_aliases(tmp_path: Path):
    """RP3: alternate skill-section labels map onto the canonical buckets so a
    resume that does not use Casey's exact headings still populates them."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Frameworks: React, Vue")  # -> Core
    doc.add_paragraph("Databases: Postgres, Redis")  # -> Data & DevOps
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert "React" in facts.skills_core
    assert "Postgres" in facts.skills_data_devops
    assert warnings == []


def test_compound_skill_labels_map_to_buckets(tmp_path: Path):
    """A9: compound labels must not drop their whole bucket.

    Regression for 2026-07-24, when a reformatted baseline used 'Languages &
    Frameworks' and 'AI & Automation'. Neither was an alias, so both buckets
    were dropped silently — verified.json reported zero core and zero AI
    skills, and the fabrication guard then rejected real skills as unverified.
    Both '&' and 'and' spellings are covered.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Languages & Frameworks: TypeScript, Next.js")
    doc.add_paragraph("AI and Automation: Claude API, Ollama")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert warnings == []
    assert "TypeScript" in facts.skills_core
    assert "Next.js" in facts.skills_core
    assert "Claude API" in facts.skills_ai
    assert "Ollama" in facts.skills_ai


def test_aspirational_labels_land_in_familiar(tmp_path: Path):
    """An aspirational row must NOT become claimable production skill.

    Regression for 2026-08-24: a lane-tailored baseline carried
    "Currently developing: RAG pipelines, embeddings and vector search, Azure".
    No Familiar keyword matched, so inference fell through to Core's "develop"
    stem and promoted skills the candidate explicitly said they were still
    learning — the one bucket error `tailoring-rules.md` treats as fabrication,
    which would then let the tailor write "Azure" onto a resume as experience.

    Matching is on the LABEL, so it holds for any lane: the items differ
    (RAG/Azure for AI, Google Ads/GA4 for programmatic) and are never
    enumerated in the parser.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Engineering: TypeScript, Python")
    doc.add_paragraph("Currently developing: RAG pipelines, Azure")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert warnings == []
    assert facts.skills_familiar == ["RAG pipelines", "Azure"]
    assert "Azure" not in facts.skills_core
    assert "RAG pipelines" not in facts.skills_core


def test_development_labels_stay_core(tmp_path: Path):
    """The Familiar keywords must not swallow real development rows.

    Familiar is tested before Core, so a "develop" stem in the Familiar
    keyword set would demote "Software Development" — inverting the bug above
    and hiding real production skill from the tailor.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Software Development: Python, Go")
    doc.add_paragraph("Web Development: React, Next.js")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert warnings == []
    assert facts.skills_familiar == []
    for skill in ("Python", "Go", "React", "Next.js"):
        assert skill in facts.skills_core


def test_technical_projects_heading_is_recognized(tmp_path: Path):
    """A9: an unaliased projects heading is absorbed *silently*.

    Regression for 2026-07-24: the resume used 'TECHNICAL PROJECTS', which was
    not an alias, so the heading and every project line became bullets on the
    preceding role — with **zero warnings**, so the convert-resume guard could
    not catch it. Silent absorption is the worst failure mode here.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Dev | Acme   Jan 2024 – Present")
    doc.add_paragraph("Shipped a thing.")
    doc.add_paragraph("TECHNICAL PROJECTS")
    doc.add_paragraph("Widget | github.com/jane/widget")
    doc.add_paragraph("Stack: Python, SQLite")
    doc.add_paragraph("Built the widget.")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert warnings == []
    assert len(facts.projects) == 1
    assert facts.projects[0].name == "Widget"
    # The role must NOT have swallowed the projects section.
    assert facts.work_history[0].bullets == ["Shipped a thing."]


def test_em_dash_project_header_parses(tmp_path: Path):
    """A9: `Name — Description — url` headers, not just `Name | url`."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("TECHNICAL PROJECTS")
    doc.add_paragraph("Widget — A Small Tool  —  github.com/jane/widget")
    doc.add_paragraph("Stack: Python, SQLite")
    doc.add_paragraph("Built it — and shipped it to production users.")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert warnings == []
    assert len(facts.projects) == 1
    assert facts.projects[0].name == "Widget — A Small Tool"
    assert facts.projects[0].url == "github.com/jane/widget"
    # The em-dash inside a prose bullet must not read as a header.
    assert facts.projects[0].bullets == [
        "Built it — and shipped it to production users."
    ]


def test_classifies_generic_certs(tmp_path: Path):
    """RP2: certs are detected by generic keywords (certified / certification /
    badge), not Casey-specific literals. The 'Associate' cert tier must not be
    mistaken for an associate degree, and Casey's Contentful + Skill Badge line
    still classifies as a cert (regression guard)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("CERTIFICATIONS & EDUCATION")
    doc.add_paragraph("AWS Certified Solutions Architect - Associate, Amazon (2024)")
    doc.add_paragraph("PMP Certification, PMI (2023)")
    doc.add_paragraph("CompTIA Security+ ce Certification (2023)")
    doc.add_paragraph("Contentful Certified Professional + Personalization Skill Badge (2025)")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert len(facts.certifications) == 4  # all four routed to certifications
    assert facts.education == []  # the AWS Associate cert did not leak to education
    assert warnings == []


def test_classifies_degrees(tmp_path: Path):
    """RP2: degree lines route to education under the generic classifier."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("CERTIFICATIONS & EDUCATION")
    doc.add_paragraph("Bachelor of Science in Computer Science, University of Toronto (2020)")
    doc.add_paragraph("Computer Programming (Advanced Diploma), George Brown College (2024)")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert len(facts.education) == 2
    assert facts.certifications == []
    assert warnings == []


def test_section_header_aliases(tmp_path: Path):
    """RP4: alternate section headings (PROFILE / SKILLS / WORK EXPERIENCE /
    EDUCATION) resolve to the canonical sections so a resume that does not use
    Casey's exact headers still parses every section."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("PROFILE")
    doc.add_paragraph("A frontend developer.")
    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Core: Python, TypeScript")
    doc.add_paragraph("WORK EXPERIENCE")
    doc.add_paragraph("Frontend Developer | Acme Inc  2022 – 2024")
    doc.add_paragraph("Built things.")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("Bachelor of Science, University of Toronto (2020)")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)
    assert facts.summary == "A frontend developer."  # PROFILE -> SUMMARY
    assert "Python" in facts.skills_core  # SKILLS -> TECHNICAL SKILLS
    assert len(facts.work_history) == 1  # WORK EXPERIENCE -> PROFESSIONAL EXPERIENCE
    assert facts.work_history[0].employer == "Acme Inc"
    assert any("Bachelor" in e for e in facts.education)  # EDUCATION -> CERTS & EDU
    assert warnings == []


def test_orphan_bullet_warns_not_raises(tmp_path: Path):
    """RP5: a bullet before any role header is warned and skipped, not raised
    (previously a fatal PipelineError that aborted the whole conversion)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Did some things before any role header.")  # orphan bullet
    doc.add_paragraph("Frontend Developer | Acme Inc  2022 – 2024")
    doc.add_paragraph("Built the thing.")
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)  # must not raise
    assert len(facts.work_history) == 1
    assert facts.work_history[0].bullets == ["Built the thing."]
    assert any("bullet before any role header" in w for w in warnings)


def test_unparseable_role_header_warns(tmp_path: Path):
    """RP5: a line with a pipe that does not match the role-header pattern (no
    parseable date) is warned and skipped, not raised. The valid role survives."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    doc.add_paragraph("Toronto, ON  |  jane@example.com")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Frontend Developer | Acme Inc  2022 – 2024")
    doc.add_paragraph("Built the thing.")
    doc.add_paragraph("Some Title | No Date Here")  # pipe but no parseable date
    path = tmp_path / "r.docx"
    doc.save(path)

    facts, warnings = parse_baseline(path)  # must not raise
    assert len(facts.work_history) == 1  # only the valid role
    assert any("unparseable role header" in w for w in warnings)


def test_role_line_accepts_parenthesized_dates():
    """2026-06 reformat: role dates wrapped in parentheses must parse, with the
    parens kept in `dates` (so the tailored resume renders them back) and the
    '(NDA)' employer suffix staying in the employer."""
    m = _ROLE_LINE_RE.match(
        "Web Developer (Contract) | Custom Jewelry Brand (Atelier Dacko) (2023 – Present)"
    )
    assert m is not None
    assert m.group("employer") == "Custom Jewelry Brand (Atelier Dacko)"
    assert m.group("dates") == "(2023 – Present)"

    m2 = _ROLE_LINE_RE.match("Web Developer (Contract) | AI Agency (NDA) (Jan 2026 – Apr 2026)")
    assert m2 is not None
    assert m2.group("employer") == "AI Agency (NDA)"
    assert m2.group("dates") == "(Jan 2026 – Apr 2026)"


def test_role_line_still_accepts_bare_dates():
    """The original bare-date format (no surrounding parens) must keep working."""
    m = _ROLE_LINE_RE.match("Sous Chef | Multiple Venues, Toronto 2015 – 2024")
    assert m is not None
    assert m.group("employer") == "Multiple Venues, Toronto"
    assert m.group("dates") == "2015 – 2024"


def test_parse_baseline_missing_file_errors(tmp_path: Path):
    with pytest.raises(PipelineError, match="not found"):
        parse_baseline(tmp_path / "nope.docx")


@pytest.mark.parametrize(
    "value,expected",
    [
        ("a, b, c", ["a", "b", "c"]),
        (
            "Shopify (Liquid, Custom Themes), HubSpot",
            ["Shopify (Liquid, Custom Themes)", "HubSpot"],
        ),
        ("foo (a, b, c), bar", ["foo (a, b, c)", "bar"]),
        ("  a  ,  b  ", ["a", "b"]),
        ("", []),
    ],
)
def test_split_skills_paren_aware(value: str, expected: list[str]):
    assert _split_skills(value) == expected


class TestProjectStackFeedsSkillsProjects:
    """A project's `Stack:` line must reach `skills_projects`.

    `skills_projects` is Core-grade (PLAN.md honesty enforcement item 1):
    technology demonstrably shipped in a public personal project. It used to be
    fed only by a "Project Stack:" row inside TECHNICAL SKILLS, so a resume
    carrying its stack on the project's own `Stack:` line parsed that line into
    `projects[].stack` and then discarded it. The bucket stayed empty, and the
    tailor could claim none of it: every item was absent from verified.json, so
    the fabrication guard read it as invented.
    """

    @staticmethod
    def _doc(tmp_path: Path, *rows: str, project_lines: tuple[str, ...] = ()) -> Path:
        from docx import Document

        doc = Document()
        doc.add_paragraph("Jane Dev")
        doc.add_paragraph("Toronto, ON  |  jane@example.com")
        doc.add_paragraph("TECHNICAL SKILLS")
        doc.add_paragraph("Core: Python, TypeScript")
        for row in rows:
            doc.add_paragraph(row)
        if project_lines:
            doc.add_paragraph("PROJECTS")
            for line in project_lines:
                doc.add_paragraph(line)
        path = tmp_path / "r.docx"
        doc.save(path)
        return path

    def test_stack_line_populates_the_bucket(self, tmp_path: Path):
        """The regression: no "Project Stack:" row anywhere, stack on the project."""
        path = self._doc(
            tmp_path,
            project_lines=(
                "Jobhunt | Job-Search CLI | github.com/jane/jobhunt",
                "Stack: asyncio, Typer, Pydantic, SQLite",
            ),
        )

        facts, _ = parse_baseline(path)

        assert facts.projects, "fixture did not parse a project"
        assert facts.projects[0].stack == ["asyncio", "Typer", "Pydantic", "SQLite"]
        # The point of the phase: the stack is claimable, not just recorded.
        assert facts.skills_projects == ["asyncio", "Typer", "Pydantic", "SQLite"]

    def test_explicit_project_stack_row_still_works(self, tmp_path: Path):
        """The previously supported shape must not regress."""
        path = self._doc(tmp_path, "Projects: FastAPI, Redis")

        facts, _ = parse_baseline(path)

        assert facts.skills_projects == ["FastAPI", "Redis"]

    def test_row_and_stack_line_are_unioned(self, tmp_path: Path):
        """Both shapes together, explicit row first."""
        path = self._doc(
            tmp_path,
            "Projects: FastAPI",
            project_lines=(
                "Jobhunt | CLI | github.com/jane/jobhunt",
                "Stack: asyncio, Typer",
            ),
        )

        facts, _ = parse_baseline(path)

        assert facts.skills_projects == ["FastAPI", "asyncio", "Typer"]

    def test_shared_library_is_listed_once(self, tmp_path: Path):
        """Two projects on the same library must not double up the bucket.

        Case-insensitive, first spelling wins, so the resume's own casing is
        what survives into verified.json.
        """
        path = self._doc(
            tmp_path,
            project_lines=(
                "Alpha | CLI | github.com/jane/alpha",
                "Stack: Typer, SQLite",
                "Beta | API | github.com/jane/beta",
                "Stack: typer, FastAPI",
            ),
        )

        facts, _ = parse_baseline(path)

        assert facts.skills_projects == ["Typer", "SQLite", "FastAPI"]

    def test_no_projects_leaves_the_bucket_empty(self, tmp_path: Path):
        """No stack anywhere must stay empty rather than inventing a bucket."""
        path = self._doc(tmp_path)

        facts, _ = parse_baseline(path)

        assert facts.skills_projects == []
