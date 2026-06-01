from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from jobhunt.pipeline.tailor import (
    TailoredCategory,
    TailoredProject,
    TailoredResume,
    TailoredRole,
)
from jobhunt.resume.render_docx import estimate_lines, fits_one_page, render


@pytest.fixture
def tailored() -> TailoredResume:
    return TailoredResume(
        summary=(
            "Full-stack developer with hands-on Shopify, HubSpot, and React work. "
            "Owns project lifecycles end-to-end. GBC diploma, Dean's List."
        ),
        skills_categories=[
            TailoredCategory("Core", ["JavaScript", "TypeScript", "React", "Next.js"]),
            TailoredCategory("CMS", ["Shopify (Liquid)", "HubSpot CMS"]),
            TailoredCategory("Familiar", ["Java", "Python"]),
        ],
        roles=[
            TailoredRole(
                title="Web Developer (Contract)",
                employer="Custom Jewelry Brand (NDA)",
                dates="2023 – Present",
                bullets=["Built 14+ page Shopify storefront.", "Shipped ring builder."],
            ),
        ],
        certifications=["Contentful Certified Professional (October 2025)"],
        education=["Computer Programming & Analysis (Advanced Diploma), GBC, April 2024"],
        coursework=["Data Structures & Algorithms", "Full-Stack Development"],
        model="qwen3.5:9b",
    )


def test_render_writes_valid_docx(tailored: TailoredResume, tmp_path: Path):
    out = render(
        tailored,
        contact_line="me@example.com | site.com",
        name="Casey Hsu",
        out_path=tmp_path / "out.docx",
    )
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Casey Hsu" in text
    assert "SUMMARY" in text
    assert "TECHNICAL SKILLS" in text
    assert "Familiar" in text
    assert "Custom Jewelry Brand" in text


def test_estimate_fits_one_page(tailored: TailoredResume):
    assert fits_one_page(tailored)
    assert estimate_lines(tailored) > 0


def test_fits_one_page_enforces_safety_margin():
    """Regression (2026-05-28): a resume estimated at exactly LINES_PER_PAGE
    rendered onto a second page in practice (the Dean's List line tipped over
    when a 6th skills category was added). fits_one_page must reserve one line
    of headroom, so an estimate equal to the raw budget is NOT a fit."""
    from jobhunt.resume.render_docx import LINES_PER_PAGE, _PAGE_SAFETY_MARGIN

    assert _PAGE_SAFETY_MARGIN >= 1
    # Build a resume padded to estimate exactly LINES_PER_PAGE, then assert it
    # is rejected (would have passed under the old zero-margin check).
    base = TailoredResume(
        summary="Full-stack developer.",
        skills_categories=[TailoredCategory("Core", ["JavaScript"])],
        roles=[TailoredRole("Dev", "Acme", "2023 – Present", ["b"])],
        certifications=[],
        education=["Diploma"],
        coursework=[],
        model="x",
    )
    pad = LINES_PER_PAGE - estimate_lines(base)
    base.roles[0].bullets.extend(["padding bullet"] * max(0, pad))
    assert estimate_lines(base) == LINES_PER_PAGE
    assert not fits_one_page(base)  # exactly-at-budget must fail with margin


def test_render_emits_single_deans_list_paragraph(tailored: TailoredResume, tmp_path: Path):
    out = render(
        tailored,
        contact_line="me@example.com",
        name="Casey Hsu",
        out_path=tmp_path / "out.docx",
    )
    doc = Document(str(out))
    deans = [p for p in doc.paragraphs if p.text.startswith("Dean")]
    assert len(deans) == 1, [p.text for p in deans]


def test_render_includes_projects_section(tailored: TailoredResume, tmp_path: Path):
    """PB4b: a tailored resume with projects renders a PROJECTS section with the
    project name, url, stack, and bullets."""
    tailored.projects = [
        TailoredProject(
            name="jobhunt",
            url="github.com/SimBuds/Jobhunt",
            stack=["Python", "Ollama", "SQLite"],
            bullets=["Built a local-first ATS CLI with a local LLM scorer."],
        )
    ]
    out = render(
        tailored,
        contact_line="me@example.com | site.com",
        name="Casey Hsu",
        out_path=tmp_path / "out.docx",
    )
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "PROJECTS" in text
    assert "jobhunt" in text
    assert "github.com/SimBuds/Jobhunt" in text
    assert "Stack: Python, Ollama, SQLite" in text
    assert "local-first ATS CLI" in text


def test_render_omits_projects_section_when_empty(tailored: TailoredResume, tmp_path: Path):
    """No projects → no PROJECTS heading (back-compat for project-less resumes)."""
    out = render(
        tailored,
        contact_line="me@example.com | site.com",
        name="Casey Hsu",
        out_path=tmp_path / "out.docx",
    )
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "PROJECTS" not in text


def test_estimate_lines_counts_projects(tailored: TailoredResume):
    """The page-fit estimate must grow when projects are added, or the shrink
    ladder would never trim them."""
    before = estimate_lines(tailored)
    tailored.projects = [
        TailoredProject("jobhunt", "github.com/SimBuds/Jobhunt", ["Python"], ["Built a CLI."]),
    ]
    assert estimate_lines(tailored) > before
