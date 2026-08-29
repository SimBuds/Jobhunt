"""Phase A9b — `convert-resume` refuses to write a partial profile.

On 2026-07-24 a reformatted baseline parsed with 12 warnings and
`convert-resume` wrote `kb/profile/` anyway. The resulting verified.json
claimed zero core skills, zero AI skills, and one role ("Sous Chef"), which
would have made the scorer grade against a hospitality profile and the
fabrication guard reject every real skill as unverified. A partial profile is
strictly worse than no profile, because nothing downstream can tell the
difference between "not in the resume" and "the parser lost it".
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pytest
from typer.testing import CliRunner

from jobhunt.commands.convert_resume_cmd import _dropped_content_warnings

runner = CliRunner()


def test_dropped_classifier_flags_lossy_warnings() -> None:
    lossy = [
        "TECHNICAL SKILLS: unrecognized skill label 'Foo', items dropped: 'A, B'",
        "PROFESSIONAL EXPERIENCE: bullet before any role header, skipped: 'x'",
        "PROFESSIONAL EXPERIENCE: unparseable role header, skipped: 'y'",
    ]
    assert _dropped_content_warnings(lossy) == lossy


def test_dropped_classifier_ignores_known_benign_warnings() -> None:
    """Warnings where the parser KEPT the content must not block a write.

    Both of these are real `parse_baseline` outputs: the cert-vs-education
    classifier falling back, and an unrecognised skill label whose items are
    still filed under Core.
    """
    benign = [
        "CERTIFICATIONS & EDUCATION: could not classify as cert or education, "
        "defaulted to education: 'Capstone: ...'",
        "TECHNICAL SKILLS: unrecognized skill label 'Wingdings', assigned to "
        "Core — add an alias if that is wrong",
    ]
    assert _dropped_content_warnings(benign) == []


def test_unrecognized_warning_kinds_block_by_default() -> None:
    """Phase A9d: the classifier fails CLOSED.

    The previous version allow-listed the words "dropped"/"skipped", so a
    future warning that discarded content while describing it differently
    would sail through the guard — the exact failure class the guard exists to
    prevent, and the same mistake the skill-label allow-list made before A9c.
    An unclassified warning must now refuse the write and get noticed.
    """
    novel = ["PROJECTS: three entries were discarded, reason unknown"]
    assert _dropped_content_warnings(novel) == novel


class TestContactLineUrls:
    """Bare domains on a contact line must resolve to `[applicant]` URLs.

    Printed resumes almost never write the `https://` scheme — there is nothing
    to click on paper. The original patterns required it, so `linkedin_url` and
    `github_url` (both in `_REQUIRED_FIELDS`) came back empty for a normally
    formatted resume, and `convert-resume` exited 1 for essentially every new
    user.
    """

    def test_bare_domains_are_extracted_and_given_a_scheme(self) -> None:
        from jobhunt.commands.convert_resume_cmd import _parse_contact_line

        got = _parse_contact_line(
            "Toronto, ON  |  416-555-0100  |  jane@example.com  |  "
            "linkedin.com/in/jane-dev  |  janedev.com  |  github.com/janedev"
        )
        assert got["linkedin_url"] == "https://linkedin.com/in/jane-dev"
        assert got["github_url"] == "https://github.com/janedev"
        assert got["portfolio_url"] == "https://janedev.com"
        assert got["phone"] == "416-555-0100"

    def test_scheme_ful_urls_still_work(self) -> None:
        from jobhunt.commands.convert_resume_cmd import _parse_contact_line

        got = _parse_contact_line(
            "jane@example.com | https://www.linkedin.com/in/jane | "
            "https://github.com/jane | https://jane.dev"
        )
        assert got["linkedin_url"] == "https://www.linkedin.com/in/jane"
        assert got["github_url"] == "https://github.com/jane"
        assert got["portfolio_url"] == "https://jane.dev"

    def test_email_domain_is_not_mistaken_for_a_portfolio(self) -> None:
        """A bare-domain pattern would otherwise match inside `x@outlook.com`."""
        from jobhunt.commands.convert_resume_cmd import _parse_contact_line

        got = _parse_contact_line("Jane Dev  |  jane@outlook.com  |  416-555-0100")
        assert "portfolio_url" not in got
        assert got["email"] == "jane@outlook.com"

    def test_dotted_library_names_are_not_portfolios(self) -> None:
        """`Node.js` / `Next.js` look like hosts; the TLD allowlist rejects them."""
        from jobhunt.commands.convert_resume_cmd import _parse_contact_line

        got = _parse_contact_line("Node.js / Next.js Developer  |  jane@example.com")
        assert "portfolio_url" not in got


class TestUserAgentBackfill:
    """The ingest User-Agent must not keep shipping a placeholder address."""

    def test_placeholder_constant_matches_the_config_default(self) -> None:
        """Pins the copy in convert_resume_cmd to the real schema default.

        The two live in different modules; if the default UA is ever reworded
        in config.py, the substring match here would silently stop firing and
        the placeholder would come back. Fail loudly instead.
        """
        from jobhunt.commands.convert_resume_cmd import (
            _DEFAULT_USER_AGENT,
            _PLACEHOLDER_CONTACT,
        )
        from jobhunt.config import IngestConfig

        assert IngestConfig().user_agent == _DEFAULT_USER_AGENT
        assert _PLACEHOLDER_CONTACT in IngestConfig().user_agent


class TestContactLineCityRegion:
    """`City, REGION` must be found wherever it sits on the line.

    The pattern was `^`-anchored, so any contact line that opened with a job
    title (the common layout) matched nothing and city/region silently kept
    their Toronto/Ontario defaults — correct for exactly one user.
    """

    @pytest.mark.parametrize(
        ("contact", "city", "region"),
        [
            # Title first, fields separated by double spaces — the layout that
            # defeated the anchored pattern.
            (
                "Full-Stack Developer  |  CMS & AI  Toronto, ON  |  a@b.io",
                "Toronto",
                "Ontario",
            ),
            # City first: what the anchored pattern was written for.
            ("Toronto, ON | 647-555-0199 | a@b.io", "Toronto", "Ontario"),
            ("Jane Dev | Vancouver, BC | jane@example.com", "Vancouver", "British Columbia"),
            # Region spelled out, and a hyphenated multi-word city.
            ("Dev | Sainte-Anne-de-Bellevue, Quebec | d@e.ca", "Sainte-Anne-de-Bellevue", "Quebec"),
            ("Engineer | Richmond Hill, Ontario | d@e.ca", "Richmond Hill", "Ontario"),
        ],
    )
    def test_city_region_is_found_anywhere_on_the_line(
        self, contact: str, city: str, region: str
    ) -> None:
        from jobhunt.commands.convert_resume_cmd import _parse_contact_line

        got = _parse_contact_line(contact)
        assert got["city"] == city
        assert got["region"] == region

    def test_leftmost_comma_does_not_win(self) -> None:
        """Anchoring on the region is what keeps a free search honest.

        An unconstrained `(...),\\s*([A-Za-z]{2,})` matches the first comma in
        the line. Here that would yield city="Developer", region="Toronto".
        """
        from jobhunt.commands.convert_resume_cmd import _parse_contact_line

        got = _parse_contact_line("Senior Developer, Toronto | Markham, ON | a@b.io")
        assert got["city"] == "Markham"
        assert got["region"] == "Ontario"

    def test_unknown_region_falls_back_rather_than_guessing(self) -> None:
        """Region allowlist is Canadian — the app is GTA/Remote-Canada scoped.

        A non-Canadian line yields nothing, leaving the config defaults in
        place, rather than recording a city paired with a region the expansion
        table cannot interpret.
        """
        from jobhunt.commands.convert_resume_cmd import _parse_contact_line

        got = _parse_contact_line("Engineer | Austin, TX | d@e.ca")
        assert "city" not in got
        assert "region" not in got


def _resume(path: Path, *, skill_label: str, role_line: str) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Dev")
    # Full contact block, split across two paragraphs like the real baseline:
    # convert-resume separately exits 1 on missing applicant fields, which
    # would mask the guard's own exit code.
    doc.add_paragraph("Toronto, ON  |  416-555-0100  |  jane@example.com")
    doc.add_paragraph(
        "linkedin.com/in/jane-dev  |  janedev.com  |  github.com/janedev"
    )
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph(f"{skill_label}: TypeScript, React")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph(role_line)
    doc.add_paragraph("Shipped a thing that worked.")
    out = path / "r.docx"
    doc.save(out)
    return out


@pytest.fixture
def env(tmp_path: Path, tmp_config_dir: Path) -> Path:
    """Scratch config so convert-resume writes into tmp, not the real kb/."""
    jh = tmp_config_dir / "jobhunt"
    jh.mkdir(parents=True, exist_ok=True)
    (jh / "config.toml").write_text(
        "[paths]\n"
        f'data_dir = "{tmp_path / "data"}"\n'
        f'db_path = "{tmp_path / "data" / "jobhunt.db"}"\n'
        f'kb_dir = "{tmp_path / "kb"}"\n'
    )
    return tmp_path


def test_convert_resume_blocks_write_when_content_dropped(env: Path) -> None:
    from jobhunt.commands import convert_resume_cmd

    docx = _resume(
        env,
        skill_label="Hobbies",  # non-skill label → items genuinely dropped
        role_line="Dev | Acme   Jan 2024 – Present",
    )
    result = runner.invoke(convert_resume_cmd.app, ["--docx", str(docx)])

    assert result.exit_code == 1
    assert "NOT written" in result.output
    assert not (env / "kb" / "profile" / "verified.json").exists()


def test_force_overrides_the_guard(env: Path) -> None:
    from jobhunt.commands import convert_resume_cmd

    docx = _resume(
        env,
        skill_label="Hobbies",
        role_line="Dev | Acme   Jan 2024 – Present",
    )
    runner.invoke(convert_resume_cmd.app, ["--docx", str(docx), "--force"])

    # Asserting on the artifact, not the exit code: `convert-resume` also exits
    # 1 on missing [applicant] fields, an unrelated pre-existing check that
    # this minimal fixture trips. The guard's contract is whether kb/profile/
    # gets written.
    assert (env / "kb" / "profile" / "verified.json").is_file()


def test_clean_parse_writes_normally(env: Path) -> None:
    from jobhunt.commands import convert_resume_cmd

    docx = _resume(
        env,
        skill_label="Languages & Frameworks",  # A9 alias
        role_line="Dev | Acme   Jan 2024 – Present",
    )
    runner.invoke(convert_resume_cmd.app, ["--docx", str(docx)])

    # See the note in test_force_overrides_the_guard on asserting the artifact
    # rather than the exit code.
    assert (env / "kb" / "profile" / "verified.json").is_file()


class TestBucketRegressions:
    """A REQUIRED skills row that vanishes from the resume must block the write.

    `parse_baseline` only warns about rows it saw and could not place, so an
    absent row produces no warning at all and `_dropped_content_warnings` has
    nothing to classify. For a required bucket that silence is dangerous: an
    emptied `skills_data_devops` is not a smaller profile but a wrong one, since
    the tailor's fabrication guard then rejects Docker/AWS/Postgres on every
    future job.

    Familiar is the exception (`_OPTIONAL_BUCKETS`). Dropping a
    "Familiar"/"Exposure" row narrows what the candidate CLAIMS, which is a
    legitimate authoring choice, and every consumer already reads the empty
    bucket as "no Familiar skills": `_complete_familiar_bucket` early-returns,
    `tailor.md` rule 10 omits the category, and `_all_matched_are_familiar`
    returns False so the Familiar-only cap never fires. It advises instead.
    """

    @staticmethod
    def _facts(**overrides: object) -> object:
        from jobhunt.resume.parse_docx import VerifiedFacts

        base: dict[str, Any] = {
            "name": "Jane Dev",
            "contact_line": "jane@example.com",
            "summary": "Dev.",
            "skills_core": ["TypeScript"],
            "skills_cms": ["Shopify"],
            "skills_data_devops": ["Docker"],
            "skills_ai": ["Claude API"],
            "skills_projects": ["FastAPI"],
            "skills_familiar": ["Java"],
            "work_history": [],
            "certifications": [],
            "education": [],
            "coursework_baseline": [],
        }
        base.update(overrides)  # type: ignore[arg-type]
        return VerifiedFacts(**base)  # type: ignore[arg-type]

    def test_emptied_required_bucket_blocks(self, tmp_path: Path) -> None:
        from jobhunt.commands.convert_resume_cmd import _bucket_regressions

        snapshot = tmp_path / "verified.json"
        snapshot.write_text(json.dumps({"skills_data_devops": ["Docker", "AWS"]}))

        blocking, advisory = _bucket_regressions(
            self._facts(skills_data_devops=[]), snapshot
        )

        assert len(blocking) == 1
        assert "Data & DevOps" in blocking[0]
        assert "2 item(s)" in blocking[0]
        assert advisory == []

    def test_emptied_familiar_advises_instead_of_blocking(
        self, tmp_path: Path
    ) -> None:
        """Familiar is optional: report it, but let the write through."""
        from jobhunt.commands.convert_resume_cmd import _bucket_regressions

        snapshot = tmp_path / "verified.json"
        snapshot.write_text(json.dumps({"skills_familiar": ["Java", "Spring Boot"]}))

        blocking, advisory = _bucket_regressions(
            self._facts(skills_familiar=[]), snapshot
        )

        assert blocking == []
        assert len(advisory) == 1
        assert "Familiar" in advisory[0]
        assert "2 item(s)" in advisory[0]

    def test_advisory_never_classifies_as_loss(self, tmp_path: Path) -> None:
        """The optional-bucket note must not reach the fail-closed guard.

        Guards the seam directly: if advisory were ever folded back into the
        warnings list, `_dropped_content_warnings` would classify it as loss
        (it carries no `_BENIGN_WARNING_MARKERS` substring) and Familiar would
        silently start blocking again.
        """
        from jobhunt.commands.convert_resume_cmd import (
            _bucket_regressions,
            _dropped_content_warnings,
        )

        snapshot = tmp_path / "verified.json"
        snapshot.write_text(json.dumps({"skills_familiar": ["Java"]}))

        blocking, advisory = _bucket_regressions(
            self._facts(skills_familiar=[]), snapshot
        )
        assert advisory, "no advisory produced, so this proves nothing"
        assert _dropped_content_warnings(blocking) == []

    def test_regression_warning_classifies_as_loss(self, tmp_path: Path) -> None:
        """The whole point: it must reach the existing fail-closed guard.

        A regression phrased so it accidentally matched `_BENIGN_WARNING_MARKERS`
        would be filtered out and never block anything.
        """
        from jobhunt.commands.convert_resume_cmd import (
            _bucket_regressions,
            _dropped_content_warnings,
        )

        snapshot = tmp_path / "verified.json"
        snapshot.write_text(json.dumps({"skills_core": ["TypeScript", "React"]}))

        blocking, _ = _bucket_regressions(self._facts(skills_core=[]), snapshot)
        assert blocking, "no regression produced, so the classification proves nothing"
        assert _dropped_content_warnings(blocking) == blocking

    def test_no_prior_snapshot_never_blocks(self, tmp_path: Path) -> None:
        """A first run has nothing to regress against."""
        from jobhunt.commands.convert_resume_cmd import _bucket_regressions

        got = _bucket_regressions(
            self._facts(skills_familiar=[], skills_core=[]),
            tmp_path / "does-not-exist.json",
        )
        assert got == ([], [])

    def test_corrupt_snapshot_never_blocks(self, tmp_path: Path) -> None:
        """A hand-edited or truncated snapshot is not evidence the resume lost
        anything, so it must not refuse an otherwise clean write."""
        from jobhunt.commands.convert_resume_cmd import _bucket_regressions

        snapshot = tmp_path / "verified.json"
        snapshot.write_text("{not json")

        assert _bucket_regressions(self._facts(skills_familiar=[]), snapshot) == ([], [])

    def test_still_empty_is_not_a_regression(self, tmp_path: Path) -> None:
        """Empty-to-empty is the steady state for buckets this resume never had.

        Without this, every run would block on a bucket the author simply does
        not use.
        """
        from jobhunt.commands.convert_resume_cmd import _bucket_regressions

        snapshot = tmp_path / "verified.json"
        snapshot.write_text(json.dumps({"skills_projects": []}))

        assert _bucket_regressions(self._facts(skills_projects=[]), snapshot) == ([], [])

    def test_populated_bucket_is_not_a_regression(self, tmp_path: Path) -> None:
        from jobhunt.commands.convert_resume_cmd import _bucket_regressions

        snapshot = tmp_path / "verified.json"
        snapshot.write_text(json.dumps({"skills_familiar": ["Java"]}))

        assert _bucket_regressions(
            self._facts(skills_familiar=["Java"]), snapshot
        ) == ([], [])

    def test_vanished_required_row_blocks_the_write_end_to_end(
        self, env: Path
    ) -> None:
        """The regression an absent row causes, driven through the real command."""
        from jobhunt.commands import convert_resume_cmd

        profile = env / "kb" / "profile"
        profile.mkdir(parents=True, exist_ok=True)
        (profile / "verified.json").write_text(
            json.dumps({"skills_data_devops": ["Docker", "AWS"]})
        )

        # This resume has no Data & DevOps row at all: its single skills row is
        # "Languages & Frameworks", which infers to Core.
        docx = _resume(
            env,
            skill_label="Languages & Frameworks",
            role_line="Dev | Acme   Jan 2024 – Present",
        )
        result = runner.invoke(convert_resume_cmd.app, ["--docx", str(docx)])

        assert result.exit_code == 1
        assert "NOT written" in result.output
        assert "Data & DevOps" in result.output
        # The prior snapshot must survive: refusing the write means keeping the
        # good profile, not replacing it with the degraded one.
        got = json.loads((profile / "verified.json").read_text())
        assert got["skills_data_devops"] == ["Docker", "AWS"]

    def test_vanished_familiar_row_still_writes(self, env: Path) -> None:
        """The 2026-08 change: an absent Familiar row must not gate the write.

        A lane-tailored baseline legitimately carries no Familiar/Exposure row.
        The write proceeds with `skills_familiar` empty, the user is told, and
        tailoring falls back to the main buckets.
        """
        from jobhunt.commands import convert_resume_cmd

        profile = env / "kb" / "profile"
        profile.mkdir(parents=True, exist_ok=True)
        (profile / "verified.json").write_text(
            json.dumps({"skills_familiar": ["Java", "Spring Boot"]})
        )

        docx = _resume(
            env,
            skill_label="Languages & Frameworks",
            role_line="Dev | Acme   Jan 2024 – Present",
        )
        result = runner.invoke(convert_resume_cmd.app, ["--docx", str(docx)])

        assert result.exit_code == 0
        assert "NOT written" not in result.output
        # Reported, not silent — the silent-empty case is what the guard exists
        # to prevent.
        assert "Familiar" in result.output
        got = json.loads((profile / "verified.json").read_text())
        assert got["skills_familiar"] == []
        assert got["skills_core"], "main buckets must still be populated"

    def test_force_still_overrides_a_regression(self, env: Path) -> None:
        from jobhunt.commands import convert_resume_cmd

        profile = env / "kb" / "profile"
        profile.mkdir(parents=True, exist_ok=True)
        (profile / "verified.json").write_text(
            json.dumps({"skills_data_devops": ["Docker"]})
        )

        docx = _resume(
            env,
            skill_label="Languages & Frameworks",
            role_line="Dev | Acme   Jan 2024 – Present",
        )
        runner.invoke(convert_resume_cmd.app, ["--docx", str(docx), "--force"])

        got = json.loads((profile / "verified.json").read_text())
        assert got["skills_data_devops"] == []
